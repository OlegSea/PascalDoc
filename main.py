from parsers.PascalABC import *
from json import load
from docx import Document

# FIXME: Не считывает опции к var в main функции
# TODO: Обработка переменной Result
# TODO: Попробовать без кавычек
# TODO: Распределить функции по файлам
# TODO: Попробовать уменьшить костыльность кода
# TODO: Улучшить универсальность регулярных выражений
# TODO: Остальные части спецификации
# TODO: Более красивое оформление спецификации
# TODO: Интерфейс

# TODO: ??? ChatGPT для написания методов

with open("config.json", "r") as f:
    config = load(f)


def replace_type(type):
    if type in config["type_replacements"]:
        return config["type_replacements"][type]
    return type


def replace_alloc_type(alloc_type):
    if alloc_type in config["alloc_type_replacements"]:
        return config["alloc_type_replacements"][alloc_type]
    return alloc_type


data = parse_folder("./example/Lab3/")

doc = Document()
doc.add_heading("Отчет по лабораторной работе №3", 0)

# construct a table for every function or procedure in all the files in the folder, and add it to the document
# table should be like this
# | variable alloc_type | variable name | variable meaning | variable type | variable structure (should be 'простая переменная' for every one| variable scope |

for file in data["files"]:
    doc.add_heading(file["file"], level=1)
    for func in file["subroutines"]:
        doc.add_heading(func["name"], level=2)
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Класс"
        hdr_cells[1].text = "Имя"
        hdr_cells[2].text = "Смысл"
        hdr_cells[3].text = "Тип"
        hdr_cells[4].text = "Структура"
        hdr_cells[5].text = "Формат"
        hdr_cells[6].text = "Ограничение"
        for var in func["variables"]:
            row_cells = table.add_row().cells
            row_cells[0].text = replace_alloc_type(var["alloc_type"])
            row_cells[1].text = var["name"]
            row_cells[2].text = var["meaning"]
            row_cells[3].text = replace_type(var["type"])
            row_cells[4].text = "Простая переменная"
            row_cells[5].text = ""
            row_cells[6].text = ""
        doc.add_paragraph()

doc.save("report.docx")
